"""Step 1: Resume Extractor - Extracts raw text from PDF/DOCX files."""

import uuid
from dataclasses import dataclass
from pathlib import Path
from typing import List

import fitz
import pytesseract
from docx import Document
from PIL import Image


@dataclass
class ExtractedResume:
  """Represents an extracted resume with raw text content."""

  engineer_id: str
  filename: str
  raw_content: str


class ResumeExtractor:
  """Extracts text content from PDF and DOCX resume files."""

  SUPPORTED = {".pdf", ".docx"}  # Removed .doc (Windows-only)

  def __init__(self, id_prefix: str = "ENG", use_uuid: bool = False) -> None:
    """Initialize the extractor.

    Args:
      id_prefix: Prefix for sequential engineer IDs.
      use_uuid: Whether to use UUIDs instead of sequential IDs.
    """
    self.id_prefix = id_prefix
    self.use_uuid = use_uuid
    self._counter = 0

  def _generate_id(self) -> str:
    """Generate a unique engineer ID."""
    if self.use_uuid:
      return str(uuid.uuid4())[:8].upper()
    self._counter += 1
    return f"{self.id_prefix}-{self._counter:03d}"

  def extract(self, file_path: str) -> str:
    """Extract text from a single resume file.

    Args:
      file_path: Path to the resume file.

    Returns:
      Extracted text content.

    Raises:
      ValueError: If file format is not supported.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in self.SUPPORTED:
      msg = f"Unsupported format: {ext}. Supported: {self.SUPPORTED}"
      raise ValueError(msg)

    if ext == ".pdf":
      return self._extract_pdf(file_path)
    return self._extract_docx(file_path)

  def _extract_pdf(self, path: str) -> str:
    """Extract text from PDF file using PyMuPDF."""
    text_parts = []
    with fitz.open(path) as doc:
      for page in doc:
        text_parts.append(page.get_text())

    text = "\n".join(text_parts).strip()

    # Use OCR for scanned PDFs with minimal text
    if len(text) < 100:
      return self._extract_pdf_ocr(path)

    return text

  def _extract_pdf_ocr(self, path: str) -> str:
    """Extract text from scanned PDF using OCR."""
    text_parts = []
    with fitz.open(path) as doc:
      for page in doc:
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        text_parts.append(pytesseract.image_to_string(img))

    return "\n".join(text_parts).strip()

  def _extract_docx(self, path: str) -> str:
    """Extract text from DOCX file."""
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]

    for table in doc.tables:
      for row in table.rows:
        for cell in row.cells:
          parts.append(cell.text)

    return "\n".join(parts).strip()

  def extract_directory(self, directory: str) -> List[ExtractedResume]:
    """Extract text from all supported files in a directory.

    Args:
      directory: Path to directory containing resume files.

    Returns:
      List of extracted resume objects.
    """
    dir_path = Path(directory)
    results = []
    files = [f for f in dir_path.iterdir() if f.suffix.lower() in self.SUPPORTED]

    if not files:
      return results

    for file_path in files:
      try:
        text = self.extract(str(file_path))
        results.append(
          ExtractedResume(
            engineer_id=self._generate_id(),
            filename=file_path.name,
            raw_content=text,
          )
        )
      except Exception:  # noqa: S110
        continue

    return results
